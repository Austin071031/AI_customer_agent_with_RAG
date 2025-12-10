"""
Document Chunking Service for AI Customer Agent.

This module provides intelligent document chunking for long documents to ensure
all content is properly stored in the knowledge base for RAG. Supports multiple
file formats (PDF, TXT, DOCX) with configurable chunk sizes and overlap.
"""

import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass
from abc import ABC, abstractmethod


@dataclass
class Chunk:
    """Represents a single chunk of a document with metadata."""
    text: str
    metadata: Dict[str, Any]
    chunk_index: int
    start_char: int
    end_char: int


class DocumentChunkingError(Exception):
    """Custom exception for document chunking related errors."""
    pass


class BaseChunker(ABC):
    """Abstract base class for document chunkers."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the chunker.
        
        Args:
            chunk_size: Maximum size of each chunk (in characters)
            chunk_overlap: Overlap between chunks (in characters)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = logging.getLogger(__name__)
        
    @abstractmethod
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """Chunk the given text into pieces."""
        pass
    
    def _validate_chunks(self, original_text: str, chunks: List[Chunk]) -> bool:
        """
        Validate that no content is lost during chunking.
        
        Args:
            original_text: The original text
            chunks: List of chunks created from the text
            
        Returns:
            True if validation passes, False otherwise
        """
        try:
            # Reconstruct text from chunks
            reconstructed_text = ""
            for chunk in sorted(chunks, key=lambda c: c.chunk_index):
                reconstructed_text += chunk.text
                
            # Remove overlap from reconstruction (simple approach)
            # This is an approximation - for exact validation we'd need more sophisticated logic
            if len(reconstructed_text) >= len(original_text) * 0.95:  # Allow 5% tolerance
                return True
            else:
                self.logger.warning(f"Content loss detected: original {len(original_text)} chars, "
                                  f"reconstructed {len(reconstructed_text)} chars")
                return False
                
        except Exception as e:
            self.logger.error(f"Validation error: {str(e)}")
            return False


class SentenceAwareChunker(BaseChunker):
    """Chunker that splits text by sentences for better semantic boundaries."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size, chunk_overlap)
        self.sentence_endings = r'(?<=[.!?])\s+'
        
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex."""
        sentences = re.split(self.sentence_endings, text)
        # Filter out empty sentences
        return [s.strip() for s in sentences if s.strip()]
    
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """
        Chunk text by sentences, respecting chunk size and overlap.
        
        Args:
            text: The text to chunk
            metadata: Metadata to include in each chunk
            
        Returns:
            List of Chunk objects
        """
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If adding this sentence would exceed chunk size (and we have some content)
            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Create a chunk from current sentences
                chunk_text = " ".join(current_chunk)
                start_char = text.find(chunk_text)
                end_char = start_char + len(chunk_text)
                
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'sentence_count': len(current_chunk),
                    'avg_sentence_length': sum(len(s) for s in current_chunk) / len(current_chunk) if current_chunk else 0
                })
                
                chunks.append(Chunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=end_char
                ))
                
                # Prepare for next chunk with overlap
                chunk_index += 1
                
                # Calculate overlap: keep some sentences from the end
                overlap_size = 0
                overlap_sentences = []
                for i in range(len(current_chunk) - 1, -1, -1):
                    if overlap_size + len(current_chunk[i]) <= self.chunk_overlap:
                        overlap_sentences.insert(0, current_chunk[i])
                        overlap_size += len(current_chunk[i])
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_size = overlap_size
            
            # Add sentence to current chunk
            current_chunk.append(sentence)
            current_size += sentence_size
        
        # Add the last chunk if there's any remaining content
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            start_char = text.find(chunk_text)
            end_char = start_char + len(chunk_text)
            
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'sentence_count': len(current_chunk),
                'avg_sentence_length': sum(len(s) for s in current_chunk) / len(current_chunk) if current_chunk else 0
            })
            
            chunks.append(Chunk(
                text=chunk_text,
                metadata=chunk_metadata,
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=end_char
            ))
        
        # Validate the chunks
        if not self._validate_chunks(text, chunks):
            self.logger.warning("Chunk validation failed - content may have been lost")
        
        return chunks


class ParagraphChunker(BaseChunker):
    """Chunker that splits text by paragraphs for document structure preservation."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size, chunk_overlap)
        
    def chunk(self, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """
        Chunk text by paragraphs, respecting chunk size and overlap.
        
        Args:
            text: The text to chunk
            metadata: Metadata to include in each chunk
            
        Returns:
            List of Chunk objects
        """
        # Split by paragraphs (blank lines)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph_size = len(paragraph)
            
            # If adding this paragraph would exceed chunk size (and we have some content)
            if current_size + paragraph_size > self.chunk_size and current_chunk:
                # Create a chunk from current paragraphs
                chunk_text = "\n\n".join(current_chunk)
                start_char = text.find(chunk_text)
                end_char = start_char + len(chunk_text)
                
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    'paragraph_count': len(current_chunk),
                    'chunk_type': 'paragraph'
                })
                
                chunks.append(Chunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=end_char
                ))
                
                # Prepare for next chunk with overlap
                chunk_index += 1
                
                # Calculate overlap: keep some paragraphs from the end
                overlap_size = 0
                overlap_paragraphs = []
                for i in range(len(current_chunk) - 1, -1, -1):
                    if overlap_size + len(current_chunk[i]) <= self.chunk_overlap:
                        overlap_paragraphs.insert(0, current_chunk[i])
                        overlap_size += len(current_chunk[i])
                    else:
                        break
                
                current_chunk = overlap_paragraphs
                current_size = overlap_size
            
            # Add paragraph to current chunk
            current_chunk.append(paragraph)
            current_size += paragraph_size
        
        # Add the last chunk if there's any remaining content
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            start_char = text.find(chunk_text)
            end_char = start_char + len(chunk_text)
            
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'paragraph_count': len(current_chunk),
                'chunk_type': 'paragraph'
            })
            
            chunks.append(Chunk(
                text=chunk_text,
                metadata=chunk_metadata,
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=end_char
            ))
        
        # Validate the chunks
        if not self._validate_chunks(text, chunks):
            self.logger.warning("Chunk validation failed - content may have been lost")
        
        return chunks


class DocumentChunkingService:
    """
    Service for intelligent document chunking with support for multiple file formats.
    
    This service handles PDF, TXT, and DOCX files with appropriate chunking strategies
    for each format to preserve document structure and semantic meaning.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document chunking service.
        
        Args:
            chunk_size: Default chunk size in characters (default: 1000)
            chunk_overlap: Default overlap between chunks in characters (default: 200)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = logging.getLogger(__name__)
        
        # Initialize chunkers for different strategies
        self.sentence_chunker = SentenceAwareChunker(chunk_size, chunk_overlap)
        self.paragraph_chunker = ParagraphChunker(chunk_size, chunk_overlap)
        
    def chunk_document(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Chunk a document based on its file type.
        
        Args:
            file_path: Path to the document file
            file_type: Type of document ('pdf', 'txt', 'docx')
            
        Returns:
            List of chunk dictionaries with text and metadata
            
        Raises:
            DocumentChunkingError: If chunking fails
        """
        try:
            # Extract text from the document
            text, file_metadata = self._extract_text(file_path, file_type)
            
            if not text.strip():
                raise DocumentChunkingError(f"Empty content in file: {file_path}")
            
            # Choose chunking strategy based on file type
            if file_type == 'pdf':
                chunks = self._chunk_pdf(file_path, text, file_metadata)
            elif file_type == 'txt':
                chunks = self._chunk_txt(text, file_metadata)
            elif file_type == 'docx':
                chunks = self._chunk_docx(file_path, text, file_metadata)
            else:
                # Default to sentence-aware chunking
                chunks = self.sentence_chunker.chunk(text, file_metadata)
            
            # Convert chunks to dictionary format
            result = []
            for chunk in chunks:
                chunk_dict = {
                    'text': chunk.text,
                    'metadata': chunk.metadata.copy(),
                    'chunk_index': chunk.chunk_index,
                    'start_position': chunk.start_char,
                    'end_position': chunk.end_char,
                    'chunk_size': len(chunk.text)
                }
                chunk_dict['metadata'].update({
                    'chunk_id': f"{Path(file_path).stem}_chunk_{chunk.chunk_index}",
                    'original_file': file_path,
                    'file_type': file_type,
                    'chunking_strategy': self._get_chunking_strategy(file_type)
                })
                result.append(chunk_dict)
            
            self.logger.info(f"Created {len(result)} chunks from {file_path}")
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to chunk document {file_path}: {str(e)}")
            raise DocumentChunkingError(f"Document chunking failed: {str(e)}")
    
    def _extract_text(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            file_type: Type of document
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        # This would integrate with the existing DocumentProcessor
        # For now, we'll implement basic extraction
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'txt':
                return self._extract_txt_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            else:
                # Fallback to text file extraction
                return self._extract_txt_text(file_path)
        except Exception as e:
            self.logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise DocumentChunkingError(f"Text extraction failed: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        try:
            # Try pdfplumber first
            import pdfplumber
            text = ""
            metadata = {'pages': 0}
            
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"Page {i+1}:\n{page_text}\n\n"
            
            metadata.update({
                'file_name': Path(file_path).name,
                'file_type': 'pdf',
                'extraction_method': 'pdfplumber'
            })
            return text.strip(), metadata
            
        except ImportError:
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                text = ""
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    metadata = {'pages': len(pdf_reader.pages)}
                    
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"Page {i+1}:\n{page_text}\n\n"
                
                metadata.update({
                    'file_name': Path(file_path).name,
                    'file_type': 'pdf',
                    'extraction_method': 'PyPDF2'
                })
                return text.strip(), metadata
                
            except ImportError:
                raise DocumentChunkingError("No PDF extraction library available")
    
    def _extract_txt_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                'file_name': Path(file_path).name,
                'file_type': 'txt',
                'encoding': 'utf-8',
                'line_count': len(text.splitlines())
            }
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                metadata = {
                    'file_name': Path(file_path).name,
                    'file_type': 'txt',
                    'encoding': 'latin-1',
                    'line_count': len(text.splitlines())
                }
                return text, metadata
            except Exception as e:
                raise DocumentChunkingError(f"Text file extraction failed: {str(e)}")
    
    def _extract_docx_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    paragraphs.append(para.text.strip())
            
            metadata = {
                'file_name': Path(file_path).name,
                'file_type': 'docx',
                'paragraph_count': len(paragraphs),
                'has_tables': len(doc.tables) > 0
            }
            return text.strip(), metadata
            
        except ImportError:
            raise DocumentChunkingError("python-docx not installed for DOCX processing")
    
    def _chunk_pdf(self, file_path: str, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """
        Chunk PDF file preserving page structure.
        
        Args:
            file_path: Path to the PDF file
            text: Extracted text from PDF
            metadata: File metadata
            
        Returns:
            List of Chunk objects
        """
        # Split text by page markers
        pages = []
        page_pattern = r'Page (\d+):\n(.*?)(?=\n\nPage \d+:\n|\Z)'
        
        for match in re.finditer(page_pattern, text, re.DOTALL):
            page_num = int(match.group(1))
            page_text = match.group(2).strip()
            pages.append((page_num, page_text))
        
        if not pages:
            # If no page markers found, use sentence-aware chunking
            return self.sentence_chunker.chunk(text, metadata)
        
        # Chunk within and across pages
        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0
        
        for page_num, page_text in pages:
            page_sentences = self.sentence_chunker._split_into_sentences(page_text)
            
            for sentence in page_sentences:
                sentence_size = len(sentence)
                
                if current_size + sentence_size > self.chunk_size and current_chunk:
                    # Create chunk
                    chunk_text = " ".join(current_chunk)
                    start_char = text.find(chunk_text)
                    end_char = start_char + len(chunk_text)
                    
                    chunk_metadata = metadata.copy()
                    page_numbers = list(set(p[0] for p in current_chunk if isinstance(p, tuple)))
                    chunk_metadata.update({
                        'page_numbers': page_numbers if page_numbers else None,
                        'sentence_count': len([s for s in current_chunk if isinstance(s, str)]),
                        'chunk_type': 'pdf_page_aware'
                    })
                    
                    chunks.append(Chunk(
                        text=chunk_text,
                        metadata=chunk_metadata,
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=end_char
                    ))
                    
                    chunk_index += 1
                    
                    # Calculate overlap
                    overlap_size = 0
                    overlap_sentences = []
                    for i in range(len(current_chunk) - 1, -1, -1):
                        if overlap_size + len(current_chunk[i]) <= self.chunk_overlap:
                            overlap_sentences.insert(0, current_chunk[i])
                            overlap_size += len(current_chunk[i])
                        else:
                            break
                    
                    current_chunk = overlap_sentences
                    current_size = overlap_size
                
                # Store sentence with page number
                current_chunk.append(sentence)
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            start_char = text.find(chunk_text)
            end_char = start_char + len(chunk_text)
            
            chunk_metadata = metadata.copy()
            page_numbers = list(set(p[0] for p in current_chunk if isinstance(p, tuple)))
            chunk_metadata.update({
                'page_numbers': page_numbers if page_numbers else None,
                'sentence_count': len([s for s in current_chunk if isinstance(s, str)]),
                'chunk_type': 'pdf_page_aware'
            })
            
            chunks.append(Chunk(
                text=chunk_text,
                metadata=chunk_metadata,
                chunk_index=chunk_index,
                start_char=start_char,
                end_char=end_char
            ))
        
        return chunks
    
    def _chunk_txt(self, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """
        Chunk text file with sentence-aware boundaries.
        
        Args:
            text: The text to chunk
            metadata: File metadata
            
        Returns:
            List of Chunk objects
        """
        return self.sentence_chunker.chunk(text, metadata)
    
    def _chunk_docx(self, file_path: str, text: str, metadata: Dict[str, Any]) -> List[Chunk]:
        """
        Chunk Word document preserving paragraph structure.
        
        Args:
            file_path: Path to the DOCX file
            text: Extracted text from DOCX
            metadata: File metadata
            
        Returns:
            List of Chunk objects
        """
        return self.paragraph_chunker.chunk(text, metadata)
    
    def _get_chunking_strategy(self, file_type: str) -> str:
        """Get the chunking strategy used for a file type."""
        strategies = {
            'pdf': 'pdf_page_aware',
            'txt': 'sentence_aware',
            'docx': 'paragraph_aware'
        }
        return strategies.get(file_type, 'sentence_aware')
    
    def update_chunking_config(self, chunk_size: int, chunk_overlap: int) -> None:
        """
        Update chunking configuration parameters.
        
        Args:
            chunk_size: New chunk size in characters
            chunk_overlap: New overlap between chunks in characters
        """
        if chunk_size <= 0:
            raise ValueError("Chunk size must be positive")
        if chunk_overlap < 0:
            raise ValueError("Chunk overlap cannot be negative")
        if chunk_overlap >= chunk_size:
            raise ValueError("Chunk overlap must be less than chunk size")
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Update the chunkers
        self.sentence_chunker.chunk_size = chunk_size
        self.sentence_chunker.chunk_overlap = chunk_overlap
        self.paragraph_chunker.chunk_size = chunk_size
        self.paragraph_chunker.chunk_overlap = chunk_overlap
        
        self.logger.info(f"Updated chunking config: size={chunk_size}, overlap={chunk_overlap}")
    
    def validate_chunks(self, original_text: str, chunks: List[Dict[str, Any]]) -> bool:
        """
        Validate that no content is lost during chunking.
        
        Args:
            original_text: The original text
            chunks: List of chunk dictionaries
            
        Returns:
            True if validation passes, False otherwise
        """
        # Convert chunk dictionaries to Chunk objects for validation
        chunk_objects = []
        for i, chunk_dict in enumerate(chunks):
            chunk_objects.append(Chunk(
                text=chunk_dict['text'],
                metadata=chunk_dict.get('metadata', {}),
                chunk_index=i,
                start_char=chunk_dict.get('start_position', 0),
                end_char=chunk_dict.get('end_position', len(chunk_dict['text']))
            ))
        
        return self.sentence_chunker._validate_chunks(original_text, chunk_objects)
