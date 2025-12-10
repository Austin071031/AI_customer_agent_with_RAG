"""
Enhanced Knowledge Base Manager for AI Customer Agent.

This module provides enhanced knowledge base management with file type detection
and intelligent routing: Excel files to SQLite database, other documents to vector storage.
Supports multiple file formats and GPU-accelerated embeddings.
"""

import os
import uuid
import logging
import shutil
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path

import torch
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

from ..models.chat_models import KBDocument
from ..models.excel_models import ExcelDocument, ExcelSearchQuery, ExcelSheetData
from .sqlite_database_service import SQLiteDatabaseService, SQLiteDatabaseError
from .document_chunking_service import DocumentChunkingService, DocumentChunkingError


class KnowledgeBaseError(Exception):
    """Custom exception for knowledge base related errors."""
    
    def __init__(self, message: str, error_type: Optional[str] = None):
        self.message = message
        self.error_type = error_type
        super().__init__(self.message)


class DocumentProcessor:
    """
    Handles document processing for various file formats.
    
    This class provides methods to extract text from different document formats
    including PDF, TXT, DOCX, and supports additional formats as needed.
    """
    
    def __init__(self):
        """Initialize the document processor with required libraries."""
        self.logger = logging.getLogger(__name__)
        self.supported_formats = {'.pdf', '.txt', '.docx', '.doc', '.md'}
        
    def extract_text_from_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from a file based on its format.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            KnowledgeBaseError: If file format is not supported or processing fails
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise KnowledgeBaseError(f"File not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise KnowledgeBaseError(f"Unsupported file format: {file_extension}")
            
        try:
            # Extract basic file metadata
            metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'last_modified': file_path.stat().st_mtime
            }
            
            # Process based on file type
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path), metadata
            elif file_extension == '.txt' or file_extension == '.md':
                return self._extract_text_file(file_path), metadata
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_text(file_path), metadata
            else:
                raise KnowledgeBaseError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {str(e)}")
            raise KnowledgeBaseError(f"Failed to process file {file_path}: {str(e)}")
            
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files using pdfplumber (primary) and PyPDF2 (fallback)."""
        text = ""
        
        # Method 1: Try pdfplumber first (better for complex PDFs)
        try:
            import pdfplumber
            self.logger.info(f"Attempting to extract PDF with pdfplumber: {file_path}")
            
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            self.logger.warning(f"Page {i+1} returned empty text with pdfplumber")
                    except Exception as page_error:
                        self.logger.warning(f"Error extracting text from page {i+1} with pdfplumber: {str(page_error)}")
                
            if text.strip():
                self.logger.info(f"Successfully extracted {len(pdf.pages)} pages with pdfplumber, total text length: {len(text)}")
                return text.strip()
            else:
                self.logger.warning(f"pdfplumber extracted empty text, trying PyPDF2 fallback")
        except ImportError:
            self.logger.warning("pdfplumber not installed, using PyPDF2")
        except Exception as e:
            self.logger.warning(f"pdfplumber extraction failed: {str(e)}, trying PyPDF2 fallback")
        
        # Method 2: Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader
            
            text = ""
            self.logger.info(f"Attempting to extract PDF with PyPDF2: {file_path}")
            
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                total_pages = len(pdf_reader.pages)
                self.logger.info(f"PDF has {total_pages} pages")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            self.logger.debug(f"Extracted {len(page_text)} chars from page {i+1}")
                        else:
                            self.logger.warning(f"Page {i+1} returned empty text with PyPDF2")
                    except Exception as page_error:
                        self.logger.warning(f"Error extracting text from page {i+1} with PyPDF2: {str(page_error)}")
                
            if text.strip():
                self.logger.info(f"Successfully extracted {total_pages} pages with PyPDF2, total text length: {len(text)}")
                return text.strip()
            else:
                raise KnowledgeBaseError("Both pdfplumber and PyPDF2 extracted empty text")
                
        except ImportError:
            raise KnowledgeBaseError("Neither pdfplumber nor PyPDF2 are installed for PDF processing")
        except Exception as e:
            raise KnowledgeBaseError(f"PDF processing error with both methods: {str(e)}")
            
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                raise KnowledgeBaseError(f"Text file encoding error: {str(e)}")
        except Exception as e:
            raise KnowledgeBaseError(f"Text file processing error: {str(e)}")
            
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files using python-docx."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except ImportError:
            raise KnowledgeBaseError("python-docx not installed for DOCX processing")
        except Exception as e:
            raise KnowledgeBaseError(f"DOCX processing error: {str(e)}")


class EmbeddingService:
    """
    Handles text embedding generation with GPU acceleration.
    
    This service uses sentence-transformers models to generate embeddings
    and automatically utilizes GPU when available for faster processing.
    """
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        """
        Initialize the embedding service.
        
        Args:
            model_name: Name of the sentence-transformers model to use
        """
        self.logger = logging.getLogger(__name__)
        self.model_name = model_name
        self.device = self._setup_device()
        self.model = self._load_model()
        
    def _setup_device(self) -> torch.device:
        """
        Configure PyTorch for GPU acceleration if available.
        
        Returns:
            torch.device: The device to use for computations
        """
        if torch.cuda.is_available():
            device = torch.device("cuda")
            torch.backends.cudnn.benchmark = True
            self.logger.info(f"Using GPU: {torch.cuda.get_device_name()}")
        else:
            device = torch.device("cpu")
            self.logger.info("Using CPU for embeddings")
        return device
        
    def _load_model(self) -> SentenceTransformer:
        """
        Load the sentence-transformers model.
        
        Returns:
            SentenceTransformer: Loaded model instance
            
        Raises:
            KnowledgeBaseError: If model loading fails
        """
        try:
            model = SentenceTransformer(self.model_name, device=self.device)
            self.logger.info(f"Loaded embedding model: {self.model_name}")
            return model
        except Exception as e:
            self.logger.error(f"Failed to load embedding model: {str(e)}")
            raise KnowledgeBaseError(f"Failed to load embedding model: {str(e)}")
            
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings for a list of texts.
        
        Args:
            texts: List of text strings to embed
            
        Returns:
            List of embedding vectors
            
        Raises:
            KnowledgeBaseError: If embedding generation fails
        """
        if not texts:
            return []
            
        try:
            # Use batch processing for efficiency
            embeddings = self.model.encode(
                texts,
                batch_size=32,
                show_progress_bar=False,
                convert_to_tensor=False,
                normalize_embeddings=True
            )
            
            # Convert numpy arrays to lists
            if hasattr(embeddings, 'tolist'):
                return embeddings.tolist()
            else:
                return [embedding.tolist() for embedding in embeddings]
                
        except Exception as e:
            self.logger.error(f"Embedding generation failed: {str(e)}")
            raise KnowledgeBaseError(f"Embedding generation failed: {str(e)}")
            
    def get_model_info(self) -> Dict[str, Any]:
        """Get information about the embedding model and device."""
        return {
            "model_name": self.model_name,
            "device": str(self.device),
            "embedding_dimension": self.model.get_sentence_embedding_dimension(),
            "gpu_available": torch.cuda.is_available()
        }


class EnhancedKnowledgeBaseManager:
    """
    Enhanced Knowledge Base Manager with file type detection and routing.
    
    This class provides intelligent file routing: Excel files to SQLite database,
    other documents to vector storage. Supports multiple file formats and
    GPU-accelerated embeddings.
    """
    
    def __init__(self, persist_directory: str = "./knowledge_base/chroma_db",
                 sqlite_db_path: str = "./excel_database.db",
                 chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the enhanced knowledge base manager.
        
        Args:
            persist_directory: Directory to persist ChromaDB data
            sqlite_db_path: Path to SQLite database for Excel files
            chunk_size: Size of document chunks (in characters)
            chunk_overlap: Overlap between document chunks (in characters)
        """
        self.logger = logging.getLogger(__name__)
        self.persist_directory = persist_directory
        self.document_processor = DocumentProcessor()
        self.embedding_service = EmbeddingService()
        self.sqlite_service = SQLiteDatabaseService(sqlite_db_path)
        self.document_chunking_service = DocumentChunkingService(
            chunk_size=chunk_size, 
            chunk_overlap=chunk_overlap
        )
        self.vector_store = self._initialize_vector_store()
        
    def _initialize_vector_store(self) -> chromadb.Collection:
        """
        Initialize the ChromaDB vector store.
        
        Returns:
            chromadb.Collection: ChromaDB collection for document storage
            
        Raises:
            KnowledgeBaseError: If vector store initialization fails
        """
        try:
            # Create persistence directory if it doesn't exist
            os.makedirs(self.persist_directory, exist_ok=True)
            
            # Initialize ChromaDB client
            client = chromadb.PersistentClient(
                path=self.persist_directory,
                settings=Settings(anonymized_telemetry=False)
            )
            
            # Create or get collection with cosine distance metric for normalized embeddings
            collection = client.get_or_create_collection(
                name="knowledge_base",
                metadata={
                    "description": "AI Customer Agent Knowledge Base",
                    "hnsw:space": "cosine"  # Use cosine distance for normalized embeddings
                }
            )
            
            self.logger.info(f"Initialized vector store at: {self.persist_directory} with cosine distance metric")
            return collection
            
        except Exception as e:
            self.logger.error(f"Failed to initialize vector store: {str(e)}")
            raise KnowledgeBaseError(f"Vector store initialization failed: {str(e)}")
            
    def _detect_file_type(self, file_path: str) -> str:
        """
        Detect file type and determine storage method.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type category: 'excel' or 'document'
        """
        file_extension = Path(file_path).suffix.lower()
        excel_extensions = {'.xlsx', '.xls', '.xlsm', '.xlsb'}
        
        if file_extension in excel_extensions:
            return 'excel'
        else:
            return 'document'
            
    def add_documents(self, file_paths: List[str], chunk_size: int = 1000, chunk_overlap: int = 200) -> Dict[str, Any]:
        """
        Process and add documents to appropriate storage based on file type.
        
        Args:
            file_paths: List of file paths to add to the knowledge base
            chunk_size: Size of document chunks (in characters)
            chunk_overlap: Overlap between document chunks (in characters)
            
        Returns:
            Dictionary with processing results for both Excel and document files
            
        Raises:
            KnowledgeBaseError: If document processing or addition fails
        """
        if not file_paths:
            raise KnowledgeBaseError("No file paths provided")
            
        # Update document chunking service with new parameters
        self.document_chunking_service.update_chunking_config(chunk_size, chunk_overlap)
            
        results = {
            'excel_files': [],
            'documents': [],
            'errors': []
        }
        
        try:
            for file_path in file_paths:
                file_path = Path(file_path)
                if not file_path.exists():
                    error_msg = f"File not found: {file_path}"
                    results['errors'].append(error_msg)
                    self.logger.error(error_msg)
                    continue
                    
                file_type = self._detect_file_type(str(file_path))
                
                try:
                    if file_type == 'excel':
                        excel_result = self._process_excel_file(str(file_path))
                        results['excel_files'].append(excel_result)
                    else:
                        document_result = self._process_document_file(str(file_path))
                        results['documents'].append(document_result)
                        
                except Exception as e:
                    error_msg = f"Failed to process file {file_path}: {str(e)}"
                    results['errors'].append(error_msg)
                    self.logger.error(error_msg)
                    
            self.logger.info(f"Processing complete: {len(results['excel_files'])} Excel files, "
                           f"{len(results['documents'])} documents, {len(results['errors'])} errors")
            return results
            
        except Exception as e:
            self.logger.error(f"Failed to add documents: {str(e)}")
            raise KnowledgeBaseError(f"Document addition failed: {str(e)}")
            
    def _process_excel_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process Excel file and store in SQLite database with dynamic table creation.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dictionary with processing results including dynamic table information
            
        Raises:
            KnowledgeBaseError: If Excel file processing fails
        """
        try:
            file_path_obj = Path(file_path)
            file_name = file_path_obj.name
            file_size = file_path_obj.stat().st_size
            
            # Extract sheet names from Excel file
            import openpyxl
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            sheet_names = workbook.sheetnames
            workbook.close()
            
            # Prepare metadata
            metadata = {
                'file_name': file_name,
                'file_size': file_size,
                'file_path': file_path,
                'uploaded_by': 'system',
                'description': f'Excel file: {file_name}',
                'dynamic_tables_created': True
            }
            
            # Store in SQLite database with dynamic table creation
            storage_result = self.sqlite_service.store_excel_file_with_dynamic_tables(
                file_path=file_path,
                file_name=file_name,
                file_size=file_size,
                sheet_names=sheet_names,
                metadata=metadata
            )
            
            result = {
                'file_id': storage_result['file_id'],
                'file_name': file_name,
                'file_size': file_size,
                'sheet_names': sheet_names,
                'storage_type': 'sqlite',
                'status': 'success',
                'tables_created': storage_result['tables_created'],
                'total_sheets': storage_result['total_sheets']
            }
            
            self.logger.info(f"Successfully processed Excel file with dynamic tables: {file_name}")
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to process Excel file {file_path}: {str(e)}")
            raise KnowledgeBaseError(f"Excel file processing failed: {str(e)}")
            
    def _clean_metadata(self, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Clean metadata by converting non-scalar values to strings and removing None values.
        
        Args:
            metadata: Original metadata dictionary
            
        Returns:
            Cleaned metadata dictionary with only scalar values, no None or empty lists
        """
        cleaned = {}
        for key, value in metadata.items():
            # Skip None values entirely
            if value is None:
                continue
                
            if isinstance(value, list):
                # Skip empty lists
                if not value:
                    continue
                # Convert non-empty list to string
                cleaned[key] = str(value)
            elif isinstance(value, dict):
                # Recursively clean nested dictionaries
                nested_cleaned = self._clean_metadata(value)
                # Only add if the nested dict is not empty
                if nested_cleaned:
                    cleaned[key] = nested_cleaned
            elif isinstance(value, (str, int, float, bool)):
                # Keep scalar values
                cleaned[key] = value
            else:
                # Convert any other type to string
                cleaned[key] = str(value)
        return cleaned

    def _process_document_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process document file with intelligent chunking and store in vector database.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with processing results including chunk information
        """
        try:
            file_type = Path(file_path).suffix.lower()
            
            # Determine file type for chunking service
            chunking_file_type = file_type[1:]  # Remove leading dot
            if chunking_file_type == 'doc':
                chunking_file_type = 'docx'  # Treat .doc as .docx for chunking
            
            # Chunk the document using the document chunking service
            chunks = self.document_chunking_service.chunk_document(
                file_path=file_path,
                file_type=chunking_file_type
            )
            
            if not chunks:
                raise KnowledgeBaseError(f"No chunks generated for file: {file_path}")
            
            # Process each chunk
            chunk_ids = []
            embeddings = []
            documents = []
            metadatas = []
            
            for chunk in chunks:
                chunk_id = str(uuid.uuid4())
                chunk_ids.append(chunk_id)
                
                # Get chunk text
                chunk_text = chunk['text']
                
                # Generate embedding for the chunk
                embedding = self.embedding_service.generate_embeddings([chunk_text])[0]
                embeddings.append(embedding)
                documents.append(chunk_text)
                
                # Prepare metadata for the chunk
                chunk_metadata = {
                    "file_path": file_path,
                    "file_type": file_type,
                    "file_name": Path(file_path).name,
                    "file_size": Path(file_path).stat().st_size,
                    "chunk_id": chunk_id,
                    "original_file": file_path,
                    "chunk_index": chunk.get('chunk_index', 0),
                    "chunk_size": len(chunk_text),
                    "start_position": chunk.get('start_position', 0),
                    "end_position": chunk.get('end_position', len(chunk_text)),
                    "chunking_strategy": chunk.get('metadata', {}).get('chunking_strategy', 'unknown'),
                    "document_id": f"{Path(file_path).stem}_doc"
                }
                
                # Merge additional chunk metadata
                if 'metadata' in chunk:
                    chunk_metadata.update(chunk['metadata'])
                
                # Clean the metadata to ensure all values are scalar types
                cleaned_metadata = self._clean_metadata(chunk_metadata)
                metadatas.append(cleaned_metadata)
            
            # Add all chunks to vector store in a single batch
            self.vector_store.add(
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas,
                ids=chunk_ids
            )
            
            result = {
                'file_name': Path(file_path).name,
                'file_size': Path(file_path).stat().st_size,
                'file_type': file_type,
                'storage_type': 'vector',
                'chunks_created': len(chunks),
                'chunk_ids': chunk_ids,
                'status': 'success',
                'chunking_strategy': chunks[0].get('metadata', {}).get('chunking_strategy', 'unknown') if chunks else 'unknown'
            }
            
            self.logger.info(f"Successfully processed document file with {len(chunks)} chunks: {Path(file_path).name}")
            return result
            
        except DocumentChunkingError as e:
            self.logger.error(f"Document chunking failed for {file_path}: {str(e)}")
            raise KnowledgeBaseError(f"Document chunking failed: {str(e)}")
        except Exception as e:
            self.logger.error(f"Failed to process document file {file_path}: {str(e)}")
            raise KnowledgeBaseError(f"Document file processing failed: {str(e)}")
            
    def search_similar(self, query: str, k: int = 3) -> List[KBDocument]:
        """
        Search for similar documents in the knowledge base (vector storage only).
        
        Args:
            query: Search query string
            k: Number of similar documents to return
            
        Returns:
            List of KBDocument objects sorted by similarity
            
        Raises:
            KnowledgeBaseError: If search operation fails
        """
        if not query.strip():
            raise KnowledgeBaseError("Empty query provided")
            
        try:
            # Generate embedding for the query
            query_embedding = self.embedding_service.generate_embeddings([query])[0]
            
            # Search in vector store with cosine distance metric
            results = self.vector_store.query(
                query_embeddings=[query_embedding],
                n_results=k,
                include=['documents', 'metadatas', 'distances']
            )
            
            # Convert results to KBDocument objects
            similar_documents = []
            if results['documents'] and results['documents'][0]:
                for i, (doc_content, metadata, distance) in enumerate(
                    zip(results['documents'][0], results['metadatas'][0], results['distances'][0])
                ):
                    kb_doc = KBDocument(
                        id=metadata['document_id'],
                        content=doc_content,
                        metadata=metadata,
                        file_path=metadata['file_path'],
                        file_type=metadata['file_type'],
                        embedding=None  # Not returning embeddings for search results
                    )
                    # Store similarity score in metadata
                    # For cosine distance: similarity = 1 - distance (distance is 1 - cosine_similarity)
                    # So similarity = cosine_similarity
                    kb_doc.metadata['similarity_score'] = 1 - distance
                    similar_documents.append(kb_doc)
                    
            self.logger.info(f"Found {len(similar_documents)} similar documents for query")
            return similar_documents
            
        except Exception as e:
            self.logger.error(f"Search failed: {str(e)}")
            raise KnowledgeBaseError(f"Search operation failed: {str(e)}")
            
    def search_excel_data(self, search_query: ExcelSearchQuery) -> List[Dict[str, Any]]:
        """
        Search for data within Excel files stored in SQLite database.
        
        Args:
            search_query: ExcelSearchQuery object with search parameters
            
        Returns:
            List of search results with file and location information
            
        Raises:
            KnowledgeBaseError: If search operation fails
        """
        try:
            results = self.sqlite_service.search_excel_data(
                query=search_query.query,
                file_id=search_query.file_id,
                sheet_name=search_query.sheet_name,
                column_name=search_query.column_name,
                case_sensitive=search_query.case_sensitive,
                max_results=search_query.max_results
            )
            
            self.logger.info(f"Found {len(results)} Excel data matches for query")
            return results
            
        except SQLiteDatabaseError as e:
            self.logger.error(f"Excel data search failed: {str(e)}")
            raise KnowledgeBaseError(f"Excel data search failed: {str(e)}")
            
    def list_excel_files(self) -> List[ExcelDocument]:
        """
        List all Excel files stored in SQLite database.
        
        Returns:
            List of ExcelDocument objects
        """
        try:
            return self.sqlite_service.list_excel_files()
        except SQLiteDatabaseError as e:
            self.logger.error(f"Failed to list Excel files: {str(e)}")
            raise KnowledgeBaseError(f"Failed to list Excel files: {str(e)}")
            
    def get_excel_file(self, file_id: str) -> Optional[ExcelDocument]:
        """
        Get specific Excel file details from SQLite database.
        
        Args:
            file_id: ID of the Excel file to retrieve
            
        Returns:
            ExcelDocument object if found, None otherwise
        """
        try:
            return self.sqlite_service.get_excel_file(file_id)
        except SQLiteDatabaseError as e:
            self.logger.error(f"Failed to get Excel file: {str(e)}")
            raise KnowledgeBaseError(f"Failed to get Excel file: {str(e)}")
            
    def delete_excel_file(self, file_id: str) -> bool:
        """
        Delete Excel file from SQLite database.
        
        Args:
            file_id: ID of the Excel file to delete
            
        Returns:
            True if deletion was successful, False otherwise
        """
        try:
            return self.sqlite_service.delete_excel_file(file_id)
        except SQLiteDatabaseError as e:
            self.logger.error(f"Failed to delete Excel file: {str(e)}")
            raise KnowledgeBaseError(f"Failed to delete Excel file: {str(e)}")
            
    def get_sheet_data(self, file_id: str, sheet_name: Optional[str] = None) -> List[ExcelSheetData]:
        """
        Get sheet data for an Excel file.
        
        Args:
            file_id: ID of the Excel file
            sheet_name: Optional specific sheet name
            
        Returns:
            List of ExcelSheetData objects
        """
        try:
            return self.sqlite_service.get_sheet_data(file_id, sheet_name)
        except SQLiteDatabaseError as e:
            self.logger.error(f"Failed to get sheet data: {str(e)}")
            raise KnowledgeBaseError(f"Failed to get sheet data: {str(e)}")
            
    def clear_knowledge_base(self) -> bool:
        """
        Clear all documents from both storage systems.
        
        Returns:
            True if both storage systems were cleared successfully, False otherwise
        """
        vector_cleared = self.clear_vector_store()
        sqlite_cleared = self.clear_sqlite_database()
            
        # Return success only if both were cleared
        if vector_cleared and sqlite_cleared:
            self.logger.info("Knowledge base cleared successfully (both vector store and SQLite database)")
            return True
        else:
            self.logger.warning(f"Knowledge base partially cleared: vector_cleared={vector_cleared}, sqlite_cleared={sqlite_cleared}")
            return False
            
    def clear_vector_store(self) -> bool:
        """
        Clear only the vector store (ChromaDB) for non-Excel documents.
        
        Returns:
            True if vector store was cleared successfully, False otherwise
        """
        try:
            # Clear vector store - try multiple methods
            try:
                # Method 1: Delete collection via client
                client = chromadb.PersistentClient(
                    path=self.persist_directory,
                    settings=Settings(anonymized_telemetry=False)
                )
                client.delete_collection("knowledge_base")
                self.logger.info("Vector store collection deleted successfully")
            except Exception as e:
                self.logger.warning(f"Failed to delete collection via client: {str(e)}")
                
                # Method 2: Delete the entire persistence directory and recreate
                try:
                    if os.path.exists(self.persist_directory):
                        self.logger.info(f"Deleting entire vector store directory: {self.persist_directory}")
                        shutil.rmtree(self.persist_directory)
                        self.logger.info("Vector store directory deleted successfully")
                except Exception as e2:
                    self.logger.error(f"Failed to delete vector store directory: {str(e2)}")
                    raise
            
            # Reinitialize vector store
            self.vector_store = self._initialize_vector_store()
            self.logger.info("Vector store cleared and reinitialized successfully")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to clear vector store: {str(e)}")
            return False
            
    def clear_sqlite_database(self) -> bool:
        """
        Clear only the SQLite database for Excel files.
        
        Returns:
            True if SQLite database was cleared successfully, False otherwise
        """
        try:
            # Clear SQLite database
            clear_result = self.sqlite_service.clear_database()
            sqlite_cleared = clear_result.get('success', False)
            if sqlite_cleared:
                self.logger.info("SQLite database cleared successfully")
                return True
            else:
                self.logger.error(f"Failed to clear SQLite database: {clear_result.get('error', 'Unknown error')}")
                return False
            
        except Exception as e:
            self.logger.error(f"Failed to clear SQLite database: {str(e)}")
            return False
            
    def get_knowledge_base_info(self) -> Dict[str, Any]:
        """
        Get comprehensive information about the knowledge base.
        
        Returns:
            Dictionary containing knowledge base statistics and information
        """
        try:
            # Get vector store info
            vector_count = self.vector_store.count()
            vector_info = {
                "document_count": vector_count,
                "persist_directory": self.persist_directory,
                "embedding_model": self.embedding_service.get_model_info(),
                "supported_formats": list(self.document_processor.supported_formats)
            }
            
            # Get SQLite database info
            sqlite_info = self.sqlite_service.get_database_info()
            
            return {
                "vector_store": vector_info,
                "sqlite_database": sqlite_info,
                "total_documents": vector_count + sqlite_info.get('document_count', 0),
                "health_status": self.health_check()
            }
            
        except Exception as e:
            self.logger.error(f"Failed to get knowledge base info: {str(e)}")
            return {"error": str(e)}
            
    def get_document_count(self) -> int:
        """
        Get the total number of documents in both storage systems.

        Returns:
            Total number of documents
        """
        try:
            vector_count = self.vector_store.count()
            sqlite_info = self.sqlite_service.get_database_info()
            sqlite_count = sqlite_info.get('document_count', 0)
            return vector_count + sqlite_count
        except Exception as e:
            self.logger.error(f"Failed to get document count: {str(e)}")
            return 0

    def get_statistics(self) -> Dict[str, Any]:
        """
        Get detailed statistics about the knowledge base.

        Returns:
            Dictionary containing knowledge base statistics
        """
        try:
            vector_count = self.vector_store.count()
            sqlite_info = self.sqlite_service.get_database_info()
            
            return {
                'total_documents': vector_count + sqlite_info.get('document_count', 0),
                'vector_documents': vector_count,
                'excel_documents': sqlite_info.get('document_count', 0),
                'total_sheets': sqlite_info.get('sheet_count', 0),
                'total_size_bytes': sqlite_info.get('total_size_bytes', 0),
                'embedding_model': self.embedding_service.model_name,
                'gpu_available': torch.cuda.is_available(),
                'last_updated': '2024-01-01T12:00:00Z'  # This would need to be tracked
            }
        except Exception as e:
            self.logger.error(f"Failed to get statistics: {str(e)}")
            return {
                'total_documents': 0,
                'vector_documents': 0,
                'excel_documents': 0,
                'total_sheets': 0,
                'total_size_bytes': 0,
                'embedding_model': self.embedding_service.model_name,
                'gpu_available': False,
                'last_updated': '2024-01-01T12:00:00Z'
            }

    def health_check(self) -> bool:
        """
        Perform a comprehensive health check on the knowledge base.

        Returns:
            True if knowledge base is healthy, False otherwise
        """
        try:
            # Test vector store
            vector_info = self.get_knowledge_base_info()
            embedding_info = self.embedding_service.get_model_info()
            
            # Test SQLite database
            sqlite_health = self.sqlite_service.health_check()
            
            # Check if all essential components are working
            if (isinstance(vector_info, dict) and 
                isinstance(embedding_info, dict) and 
                self.vector_store is not None and
                sqlite_health):
                return True
            return False

        except Exception:
            return False

# For backward compatibility, alias the enhanced class as KnowledgeBaseManager
KnowledgeBaseManager = EnhancedKnowledgeBaseManager
